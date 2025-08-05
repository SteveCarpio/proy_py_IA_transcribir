######################################################################################################################
# Primera vez: crear :  py -3.12 -m venv venv312
# .\venv312\Scripts\activate
# pip install pydub
# pip install vosk 
# pip install ffmpeg-python
# pip install requests
# pip freeze > requirements.txt
# https://alphacephei.com/vosk/models   # -- descargar el modelos  (por ejemplo “vosk-model-small-es-0.42”).
# Descomprime el archivo ZIP y copiar la carpeta en una ruta y ponerlo en el programa
# Agregar al path la ruta con la carpeta ..../bin del directorio zapeado...
######################################################################################################################

import os  # Importación de la librería para operaciones de sistema (manejo de archivos y directorios)
import json  # Importación de la librería para trabajar con datos en formato JSON
from vosk import Model, KaldiRecognizer  # Importación de las clases necesarias para el reconocimiento de voz
import wave  # Librería para trabajar con archivos WAV
from pydub import AudioSegment  # Librería para convertir y manipular audios
import requests  # Librería para realizar solicitudes HTTP
from docx import Document
import re

######################################################################################################################
# Funciones Varias
######################################################################################################################

# Guardar la transcripción como DOCX:
def save_docx(text, base):
    docx_transcription_path = base + "_completo.docx"
    doc = Document()
    doc.add_heading('Transcripción de la Reunión', 0)
    
    # Agregar cada línea de texto como un párrafo
    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.save(docx_transcription_path)


def save_markdown_to_docx(markdown_text, base):
    docx_summary_path = base + "_resumen.docx"
    # Crear un documento Word vacío
    doc = Document()
    doc.add_heading('Resumen Profesional de la Reunión', 0)

    # Expresiones regulares para detectar Markdown
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')  # Buscar negrita (texto entre **)
    list_pattern = re.compile(r'\* (.*?)\n')  # Buscar listas (texto que empieza con * y un espacio)
    
    # Dividir el texto por líneas
    lines = markdown_text.split('\n')
    
    for line in lines:
        if line.startswith("**"):  # Si la línea tiene negrita
            # Encontrar el texto entre ** y aplicarle negrita
            bold_text = bold_pattern.sub(r'\1', line)
            p = doc.add_paragraph(bold_text, style='Heading 1')  # Estilo de título (puedes cambiarlo)
            p.paragraph_format.space_after = 0  # Quitar el espacio después del párrafo
        elif line.startswith("*"):  # Si la línea es una lista
            p = doc.add_paragraph(line[2:], style='List Bullet')  # Agregar lista con viñetas
            p.paragraph_format.space_after = 0  # Quitar el espacio después del párrafo
        else:
            p = doc.add_paragraph(line)  # Si es solo texto, agregar normalmente
            p.paragraph_format.space_after = 0  # Quitar el espacio después del párrafo

    # Guardar el documento en el archivo .docx
    doc.save(docx_summary_path)



# Función: convert_to_wav
# Objetivo: Convertir cualquier archivo de audio a formato WAV con un solo canal y frecuencia de 16000
def convert_to_wav(audio_path):
    # Cargar el archivo de audio y convertirlo a formato WAV con las configuraciones especificadas
    sound = AudioSegment.from_file(audio_path)
    sound = sound.set_channels(1).set_frame_rate(16000) # Configura el canal mono y la frecuencia de muestreo
    wav_path = "temp.wav"
    sound.export(wav_path, format="wav") # Exportar el audio convertido
    return wav_path

# Función: transcribe
# Objetivo: Transcribir el audio (en formato WAV) a texto utilizando el modelo de Vosk
def transcribe(audio_path, model_path):
    # Convertir el audio de entrada a formato WAV
    wav_path = convert_to_wav(audio_path)
    # Cargar el modelo preentrenado de Vosk

    model = Model(model_path)
    results = [] # Lista para almacenar las transcripciones

    # Abrir el archivo WAV y preparar el reconocimiento de voz
    with wave.open(wav_path, "rb") as wf:
        rec = KaldiRecognizer(model, wf.getframerate())
        while True:
            # Leer datos del archivo WAV en bloques de 4000 frames
            data = wf.readframes(4000)
            if len(data) == 0: # Si no hay más datos, salir del bucle
                break
            if rec.AcceptWaveform(data): # Si el bloque de audio es aceptado por el reconocedor
                # Solo texto puro
                result_json = json.loads(rec.Result()) # Parsear la transcripción en formato JSON
                results.append(result_json.get("text", "")) # Guardar el texto transcrito
        # Obtener el resultado final de la transcripción
        final_json = json.loads(rec.FinalResult())
        results.append(final_json.get("text", "")) # Guardar el texto final

    # Eliminar el archivo WAV temporal
    os.remove(wav_path)
    # Unir todas las transcripciones y devolverlas como un único texto
    return "\n".join(results).strip() 

# Función: save_txt
# Objetivo: Guardar el texto transcrito en un archivo de texto (.txt)
def save_txt(text, path):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text) # Escribir el texto en el archivo especificado

# Función: resumir_ollama
# Objetivo: Generar un resumen profesional de la reunión usando un modelo de IA (por ejemplo, Ollama)
def resumir_ollama(texto, modelo, base):
    url = "http://localhost:11434/api/generate"  # URL del servidor local de Ollama para generar el resumen
    
    # Crear el prompt para pedir el resumen a la IA
    prompt = (
    "A continuación tienes la transcripción de una reunión en español, posiblemente sin puntuación ni formato."
    "Tu tarea es redactar un acta profesional de la reunión, respetando el idioma español."
    "\n\nPor favor, sigue estas instrucciones:"
    "\n- Indica la fecha, hora y lugar de la reunión (si esos datos aparecen en el texto; si no, deja un espacio para completarlos)."
    "\n- Presenta una breve descripción de los participantes, indicando nombre completo y, si es posible, su empresa o rol."
    "\n- Haz un listado breve y claro de los puntos tratados en la reunión."
    "\n- Si quedan temas pendientes o para la próxima reunión, indícalos como 'Puntos pendientes'."
    "\n- Mantén una redacción clara, profesional y estructurada."
    "\n\nTranscripción de la reunión:"
    f"\n\n{texto}\n\n"
)

    data = {
        "model": modelo,
        "prompt": prompt,
        "stream": False    # Indicador de si se desea streaming de la respuesta
    }
    # Realizar la solicitud POST al servidor de Ollama para obtener el resumen
    response = requests.post(url, json=data)
    response.raise_for_status() # Verificar si hubo algún error en la solicitud
    resumen = response.json()["response"].strip()  # Obtener el resumen del texto del audio

    # Crear archivo de texto con el resumen obtenido
    resumen_path = base + "_resumen.txt"
    save_txt(resumen, resumen_path)                # Guardar el resumen en el archivo

    return resumen, resumen_path

# Función: procesar_audio
# Objetivo: Transcribir el audio y luego generar un resumen utilizando la IA
def procesar_audio(audio_file, modelo_dir, base):

    # Llamar a la función de transcripción para obtener el texto del audio
    texto = transcribe(audio_file, modelo_dir)

    # Guardar el texto transcrito en un archivo .txt
    txt_path = base + "_completo.txt"
    save_txt(texto, txt_path)

    return texto, txt_path

######################################################################################################################
#                                              Inicio del Programa                                                   #
######################################################################################################################

if __name__ == "__main__":

    audio_file = "C:\\MisCompilados\\audios\\REUNION.mp3"
    modelo_dir  = "C:\\MisCompilados\\utils\\model\\vosk-model-es-0.42"      #  [ vosk-model-small-es-0.42 ]
    modelo_ollama = "llama3:instruct"                                        #  [ llama3:instruct | mistral ]
    base, _ = os.path.splitext(audio_file)

    # Función Procesar Audio
    texto, txt_path = procesar_audio(audio_file, modelo_dir, base)  
    print(f"\n - Transcripción guardada en: {txt_path}")
    print(f"\n{texto}")

    save_docx(texto, base)

    # Función Crea Resumen Modelo IA
    resumen, resumen_path = resumir_ollama(texto, modelo_ollama, base)
    print(f"\n\n - Resumen guardado en: {resumen_path}")
    print(f"\n{resumen}\n")
    
    #save_summary_docx(resumen, base)
    save_markdown_to_docx(resumen, base)